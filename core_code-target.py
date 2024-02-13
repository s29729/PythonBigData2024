import openai
API_KEY = ""

import whisper
import whisper_timestamped as whisper
import os
import json
from openai import OpenAI

from moviepy.editor import VideoFileClip, concatenate_videoclips


model = whisper.load_model("base")

result = model.transcribe("MSTeams video.mp4")

transcription = result["text"]
print(transcription)


client = OpenAI(
    # This is the default and can be omitted
    api_key=API_KEY
)


intrukcja = client.chat.completions.create(
    messages=[
        {
            "role": "system",
            "content": "You are a helpful assistant. Please reformat the following transcription into a clear, step-by-step instructional format."
        },
        {
            "role": "user",
            "content": transcription,
        }
    ],
    model="gpt-3.5-turbo",
)
intrukcja.choices[0].message.content


audio = whisper.load_audio("MSTeams video.mp4")

model = whisper.load_model("tiny", device="cpu")

result_with_timestamps = whisper.transcribe(model, audio, language="pl")

print(json.dumps(result_with_timestamps, indent = 2, ensure_ascii = False))

moments = json.dumps(result_with_timestamps, indent = 2, ensure_ascii = False)

data = json.loads(moments)

# Extracting only the necessary fields
simplified_data = []
for segment in data.get("segments", []):
    simplified_segment = {
        "text": segment.get("text"),
        "start": segment.get("start"),
        "end": segment.get("end")
    }
    simplified_data.append(simplified_segment)

# Convert the simplified data back to JSON
simplified_moments = json.dumps(simplified_data, indent=2, ensure_ascii=False)
print(simplified_moments)

result_instruction = client.chat.completions.create(
    messages=[
        {
            "role": "system",
            "content": "You are a helpful assistant. Create a list of start moments of important moments in the transcription based on included timestamps. Make sure the content is in the format as in the following example: '- 0:06.68 - W sekcji czatów wyszukuje kasię kowalczyk,\n- 0:29.16 - Potrzebuję abyś przygotowała rapor do tryczący naszego'"
        },
        {
            "role": "user",
            "content": simplified_moments,
        }
    ],
    model="gpt-3.5-turbo",
)
result_instruction.choices[0].message.content

ts_result = client.chat.completions.create(
    messages=[
        {
            "role": "system",
            "content": "You are a helpful assistant. Create a list of start moments of important moments in the transcription based on included timestamps in HH:MM:SS format, make sure miliseconds are excluded. Here is the example of an output ['00:00:06', '00:00:16', '00:00:22']"
        },
        {
            "role": "user",
            "content": simplified_moments,
        }
    ],
    model="gpt-3.5-turbo",
)
ts_result.choices[0].message.content

# Given string of timestamps
timestamps_str = ts_result.choices[0].message.content
# Convert the string to a list of timestamps
# Using eval() to evaluate the string as a Python expression
# Note: Using eval() can be dangerous if you're working with untrusted input. In this case, it's okay since the input format is known and controlled.
timestamps_list = eval(timestamps_str)

# Verify the conversion
print(timestamps_list)

from moviepy.editor import VideoFileClip, ImageClip
import os

def take_screenshots(video_path, timestamps, output_folder):
    # Ensure the output folder exists
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Load the video file
    with VideoFileClip(video_path) as video:
        for i, timestamp in enumerate(timestamps):
            # Convert timestamp to seconds
            hours, minutes, seconds = map(int, timestamp.split(':'))
            total_seconds = hours * 3600 + minutes * 60 + seconds
# Take screenshot at the specified timestamp
            frame = video.get_frame(total_seconds)
            screenshot = ImageClip(frame).set_duration(1)
            timestamps_list = timestamp.replace(':', '-')
            output_path = os.path.join(output_folder, f"screenshot_{timestamps_list}.png")
            screenshot.save_frame(output_path)

# Example usage
video_path = "MSTeams video.mp4"
timestamps = timestamps_list
output_folder = 'screenshots'

take_screenshots(video_path, timestamps, output_folder)

ts_result = client.chat.completions.create(
    messages=[
        {
            "role": "system",
            "content": "You are a helpful assistant. Create a list of start moments of important moments in the transcription based on included timestamps in MM:SS.mS format, make sure miliseconds are included. Here is the example of an output ['0:06.68', '0:16.68', '0:23.68']"
        },
        {
            "role": "user",
            "content": simplified_moments,
        }
    ],
    model="gpt-3.5-turbo",
)
ts_result.choices[0].message.content
# Given string of timestamps
timestamps_str = ts_result.choices[0].message.content
# Convert the string to a list of timestamps
# Using eval() to evaluate the string as a Python expression
# Note: Using eval() can be dangerous if you're working with untrusted input. In this case, it's okay since the input format is known and controlled.
timestamps_list1 = eval(timestamps_str)

# Verify the conversion
print(timestamps_list1)

instruction = client.chat.completions.create(
    messages=[
        {
            "role": "system",
            "content": "You are a helpful assistant. Please reformat the following transcription into a clear, step-by-step instructional format, including timestamps. Make sure that the generated content includes only start timestamp and instruction in the following format: MM:SS.Ms - 'content', example: 00:06.68 - In the chat section, search for Kasia Kowalczyk using the search function or by finding her on the list of recent conversations."
        },
        {
            "role": "user",
            "content": result_instruction.choices[0].message.content,
        }
    ],
    model="gpt-3.5-turbo",
)

instruction_input = instruction.choices[0].message.content
print(instruction_input)

from docx import Document
from docx.shared import Inches
import os

def create_instruction_document(instruction_input, output_folder='screenshots'):
    # Initialize a new Document
    doc = Document()
    doc.add_heading('Step-by-Step Instructions', level=1)

    # Split the instruction input by line breaks to get individual instructions
    steps = instruction_input.strip().split('\n')

    for i, step in enumerate(steps):
        # Extract the timestamp and step description
        timestamp, description = step.split(' - ', 1)
        # Convert timestamp to match the screenshot filename format
        minutes, rest = timestamp.split(':')
        seconds = rest.split('.')[0]  # Ignore milliseconds
        # Padding minutes and seconds with zero if necessary
        timestamps_list = f"00-{int(minutes):02d}-{int(seconds):02d}"

        # Add instruction to the document
        doc.add_heading(f'Step {i+1}:', level=2)
        doc.add_paragraph(description)

        # Embed the corresponding screenshot
        screenshot_path = os.path.join(output_folder, f'screenshot_{timestamps_list}.png')

        if os.path.exists(screenshot_path):
            doc.add_picture(screenshot_path, width=Inches(5.0))
        else:
            doc.add_paragraph(f"(Screenshot not found for timestamp {timestamps_list})")

    # Save the document
    output_path = os.path.join(output_folder, 'instructions.docx')
    doc.save(output_path)
    return output_path


output_path = create_instruction_document(instruction_input)
print(f"Instruction document saved at: {output_path}")